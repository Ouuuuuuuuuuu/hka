import streamlit as st
import pandas as pd
import docx
import json
import requests
import io
import time
import zipfile
import os
import asyncio
import aiohttp
import nest_asyncio
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Any, Tuple
import hashlib
import pickle
from dataclasses import dataclass

# 应用 nest_asyncio 以支持在 Streamlit 中运行异步代码
nest_asyncio.apply()

# 可选依赖
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    st.error("请安装 pymupdf: pip install pymupdf")
    st.stop()

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import rarfile
    RAR_SUPPORT = True
except ImportError:
    RAR_SUPPORT = False

try:
    import py7zr
    SEVENZ_SUPPORT = True
except ImportError:
    SEVENZ_SUPPORT = False

# ============================
# 新增：文件名解码（解决乱码）
# ============================
def decode_filename(filename: str) -> str:
    """智能解码文件名，处理各种编码乱码"""
    import re
    
    # 如果已经是正常中文，直接返回
    if len(re.findall(r'[\u4e00-\u9fff]', filename)) >= 2:
        return filename
    
    # 尝试多种编码解码
    encodings = ['gbk', 'gb2312', 'gb18030', 'utf-8', 'cp437', 'latin1', 'big5']
    
    for encoding in encodings:
        try:
            if isinstance(filename, str):
                raw_bytes = filename.encode('latin1', errors='ignore')
            else:
                raw_bytes = filename
            decoded = raw_bytes.decode(encoding, errors='ignore')
            if len(re.findall(r'[\u4e00-\u9fff]', decoded)) >= 2:
                return decoded
        except:
            continue
    
    return filename

# ============================
# 新增：OCR 支持（DeepSeek 优先）
# ============================
async def ocr_recognize(file_bytes: bytes, file_ext: str, api_key: str) -> str:
    """OCR 识别，优先使用 DeepSeek"""
    import base64
    
    try:
        base64_data = base64.b64encode(file_bytes).decode('utf-8')
        mime_type = {
            'pdf': 'application/pdf', 'png': 'image/png', 'jpg': 'image/jpeg'
        }.get(file_ext.lower(), 'application/octet-stream')
        
        data_url = f"data:{mime_type};base64,{base64_data}"
        
        url = "https://api.siliconflow.cn/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "model": "deepseek-ai/DeepSeek-V3",
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": data_url}},
                    {"type": "text", "text": "请识别图片/PDF中的全部文字内容，保持原有排版格式返回。"}
                ]
            }],
            "max_tokens": 4000,
            "temperature": 0.1
        }
        
        async with aiohttp.ClientSession() as session:
            async with session.post(url, headers=headers, json=payload, timeout=60) as response:
                if response.status == 200:
                    data = await response.json()
                    return data['choices'][0]['message']['content']
    except Exception as e:
        print(f"OCR 失败: {e}")
    
    return ""

# ============================
# 配置区 - 通用化设计
# ============================
st.set_page_config(
    page_title="智能简历筛选系统", 
    layout="wide", 
    page_icon="📄",
    initial_sidebar_state="expanded"
)

# 可配置参数
DEFAULT_CONFIG = {
    "target_city": "",  # 空字符串表示不限制地域
    "max_workers": 10,  # 增加并发数
    "batch_size": 20,   # 批处理大小
    "api_timeout": 60,  # API超时时间
    "enable_cache": True,  # 启用解析缓存
    "stream_output": True,  # 流式输出结果
}

# 立即初始化所有 session_state 变量
if 'config' not in st.session_state:
    st.session_state.config = DEFAULT_CONFIG.copy()

if 'uploaded_files_queue' not in st.session_state:
    st.session_state.uploaded_files_queue = []

if 'final_results' not in st.session_state:
    st.session_state.final_results = []

if 'need_review' not in st.session_state:
    st.session_state.need_review = []

# ============================
# 缓存系统 - 提升重复处理效率
# ============================
class ResumeCache:
    """简历解析缓存系统，避免重复解析相同文件"""
    
    def __init__(self, enable_cache=True):
        # 使用临时目录，兼容 Streamlit Cloud
        import tempfile
        self.cache_dir = os.path.join(tempfile.gettempdir(), "resume_cache")
        self.enable_cache = enable_cache
        try:
            os.makedirs(self.cache_dir, exist_ok=True)
        except Exception as e:
            self.cache_dir = None
    
    def _get_cache_key(self, file_bytes: bytes) -> str:
        """基于文件内容生成缓存key"""
        return hashlib.md5(file_bytes).hexdigest()
    
    def get(self, file_bytes: bytes) -> Dict:
        """获取缓存的解析结果"""
        if not self.enable_cache or self.cache_dir is None:
            return None
        cache_key = self._get_cache_key(file_bytes)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                return None
        return None
    
    def set(self, file_bytes: bytes, result: Dict):
        """保存解析结果到缓存"""
        if not self.enable_cache or self.cache_dir is None:
            return
        cache_key = self._get_cache_key(file_bytes)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.pkl")
        try:
            with open(cache_file, 'wb') as f:
                pickle.dump(result, f)
        except Exception as e:
            pass

# 创建全局缓存实例 - 使用固定的 enable_cache
cache = ResumeCache(enable_cache=True)

# ============================
# JSON Schema - 简历结构化标准
# ============================
JSON_SCHEMA = """
{
    "basic_info": {
        "name": "姓名",
        "gender": "性别(男/女)",
        "age": "年龄(数字)",
        "phone": "电话",
        "subject": "任教学科",
        "marital_status": "婚姻状况(未婚/已婚/离异)",
        "residence": "现居地",
        "partner_location": "配偶/伴侣所在地",
        "parents_background": "父母职业背景"
    },
    "education": {
        "high_school": "高中学校",
        "high_school_tier": "高中层次(如省重点/市重点)",
        "bachelor_school": "本科学校",
        "bachelor_tier": "本科层次(如C9/985/211/双一流/普通一本/海外名校)",
        "bachelor_major": "本科专业",
        "master_school": "硕士学校",
        "master_tier": "硕士层次(如C9/985/211/双一流/普通一本/海外名校)",
        "master_major": "硕士专业",
        "study_abroad_years": "留学年限(年)",
        "exchange_experience": "是否有交换经历(是/否)"
    },
    "work_experience": {
        "current_company": "现工作单位",
        "school_tier": "学校层次(如市重点/国际学校/知名学校等)",
        "non_teaching_gap": "非教学空窗期(年)",
        "gap_explanation_valid": "空窗期解释是否合理(是/否/待定)",
        "overseas_work_years": "海外工作年限(年)",
        "management_role": "管理职务(如年级组长/教研组长/中层/主任/副校长等)",
        "head_teacher_years": "班主任年限(年)",
        "teaching_years": "教龄(年)"
    },
    "achievements": {
        "honor_titles": ["荣誉称号列表，如特级教师/学科带头人/骨干教师/优青等"],
        "teaching_competition": ["教学竞赛获奖列表"],
        "academic_results": ["学术成果列表，如论文/课题等"]
    },
    "ai_assessment": {
        "summary": "综合评语",
        "teaching_philosophy": "教学理念",
        "resume_quality_score": "简历质量评分(1-5分)",
        "career_trajectory": "职业发展轨迹(上升/平稳/波动)",
        "potential_score": "AI潜力评分(1-5分)",
        "risk_warning": "风险提示(如有)"
    }
}
"""

# ============================
# 文件解析 - 支持批量和缓存
# ============================
@st.cache_data(ttl=3600, show_spinner=False)
def extract_text_from_pdf_cached(file_bytes: bytes) -> str:
    """带缓存的PDF解析"""
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            return "\n".join([line.strip() for line in text.split('\n') if line.strip()])
    except Exception as e:
        return f"PDF解析错误: {str(e)}"

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """优先从缓存获取PDF解析结果"""
    cached = cache.get(file_bytes)
    if cached and 'pdf_text' in cached:
        return cached['pdf_text']
    
    result = extract_text_from_pdf_cached(file_bytes)
    cache.set(file_bytes, {'pdf_text': result})
    return result

def extract_text_from_docx(file_bytes: bytes, file_name: str = "") -> str:
    """DOCX/DOC解析，支持多种方式"""
    text = ""
    
    # 尝试python-docx
    if DOCX_SUPPORT and file_name.lower().endswith('.docx'):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            if text.strip():
                return text
        except:
            pass
    
    # 回退到PyMuPDF
    try:
        filetype = "doc" if file_name.lower().endswith('.doc') else "docx"
        with fitz.open(stream=file_bytes, filetype=filetype) as doc:
            for page in doc:
                text += page.get_text() + "\n"
        if text.strip():
            return text
    except:
        pass
    
    return "Word文档解析失败"

def extract_archive_files(file_bytes: bytes, file_name: str) -> List[Dict]:
    """批量解压压缩文件 - 新增文件名解码"""
    extracted_files = []
    supported_ext = ('.pdf', '.docx', '.doc')
    
    try:
        if file_name.lower().endswith('.zip'):
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                for item in zf.namelist():
                    if item.lower().endswith(supported_ext):
                        raw_name = os.path.basename(item)
                        decoded_name = decode_filename(raw_name)  # 解码文件名
                        extracted_files.append({
                            'name': decoded_name,
                            'original_name': raw_name,
                            'bytes': zf.read(item)
                        })
        
        elif file_name.lower().endswith('.rar') and RAR_SUPPORT:
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.rar') as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            with rarfile.RarFile(tmp_path) as rf:
                for item in rf.namelist():
                    if item.lower().endswith(supported_ext):
                        raw_name = os.path.basename(item)
                        decoded_name = decode_filename(raw_name)
                        extracted_files.append({
                            'name': decoded_name,
                            'original_name': raw_name,
                            'bytes': rf.read(item)
                        })
            os.unlink(tmp_path)
        
        elif file_name.lower().endswith(('.7z', '.7zip')) and SEVENZ_SUPPORT:
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.7z') as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            with py7zr.SevenZipFile(tmp_path, mode='r') as zf:
                for name, bio in zf.readall().items():
                    if name.lower().endswith(supported_ext):
                        raw_name = os.path.basename(name)
                        decoded_name = decode_filename(raw_name)
                        extracted_files.append({
                            'name': decoded_name,
                            'original_name': raw_name,
                            'bytes': bio.read()
                        })
            os.unlink(tmp_path)
            
    except Exception as e:
        st.error(f"解压失败 {file_name}: {str(e)}")
    
    return extracted_files

# ============================
# 批量文件解析 - 高效并发
# ============================
@dataclass
class ParseResult:
    """解析结果数据结构"""
    filename: str
    original_name: str  # 新增：原始文件名（乱码）
    content: str
    error: str = None
    file_size: int = 0
    parse_time: float = 0.0
    ocr_used: bool = False  # 新增：是否使用了 OCR

def parse_single_file(item: Dict, api_key: str = None, use_ocr: bool = False) -> ParseResult:
    """单文件解析，带性能统计 - 新增 OCR 支持"""
    start_time = time.time()
    try:
        file_name = item['name']
        original_name = item.get('original_name', file_name)  # 保存原始文件名
        content_bytes = item['bytes']
        file_size = len(content_bytes)
        ocr_used = False
        
        # 检查缓存
        cached = cache.get(content_bytes)
        if cached and 'parsed_result' in cached:
            result = cached['parsed_result']
            result.parse_time = time.time() - start_time
            return result
        
        text = ""
        error_msg = None
        
        if file_name.lower().endswith('.pdf'):
            text = extract_text_from_pdf(content_bytes)
        elif file_name.lower().endswith(('.docx', '.doc')):
            text = extract_text_from_docx(content_bytes, file_name)
        else:
            error_msg = "不支持的文件类型"
        
        # 新增：文本过短时尝试 OCR
        if len(text) < 50 and not error_msg and use_ocr and api_key:
            try:
                import asyncio
                loop = asyncio.get_event_loop()
                ocr_text = loop.run_until_complete(
                    ocr_recognize(content_bytes, file_name.split('.')[-1], api_key)
                )
                if ocr_text and len(ocr_text) > 50:
                    text = ocr_text
                    ocr_used = True
            except Exception as e:
                print(f"OCR 失败: {e}")
        
        if len(text) < 50 and not error_msg:
            error_msg = "提取文本过短，可能解析失败"
        
        result = ParseResult(
            filename=file_name,
            original_name=original_name,
            content=text,
            error=error_msg,
            file_size=file_size,
            parse_time=time.time() - start_time,
            ocr_used=ocr_used
        )
        
        # 存入缓存
        cache.set(content_bytes, {'parsed_result': result})
        return result
        
    except Exception as e:
        return ParseResult(
            filename=item.get('name', 'unknown'),
            original_name=item.get('original_name', item.get('name', 'unknown')),
            content="",
            error=str(e),
            parse_time=time.time() - start_time,
            ocr_used=False
        )

def parse_files_batch(uploaded_items: List[Dict], progress_callback=None, api_key: str = None, use_ocr: bool = False) -> Tuple[List[ParseResult], List[Dict]]:
    """批量文件解析，使用线程池并发处理 - 新增 OCR 支持"""
    parsed_data = []
    failed_files = []
    max_workers = st.session_state.config.get('max_workers', 10)
    
    st.write(f"🔧 启动解析: {len(uploaded_items)} 个文件, {max_workers} 个并发 worker")
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_item = {
            executor.submit(parse_single_file, item, api_key, use_ocr): item 
            for item in uploaded_items
        }
        
        for i, future in enumerate(as_completed(future_to_item)):
            try:
                result = future.result()
                item = future_to_item[future]
                if result.error:
                    # 新增：保存完整失败信息用于下载
                    failed_files.append({
                        'name': result.filename,
                        'original_name': result.original_name,
                        'error': result.error,
                        'size': result.file_size,
                        'bytes': item.get('bytes', b'')  # 保存文件字节用于打包下载
                    })
                    ocr_info = " (已尝试OCR)" if result.ocr_used else ""
                    st.write(f"❌ 解析失败: {result.filename}{ocr_info} - {result.error}")
                else:
                    parsed_data.append(result)
                    ocr_info = " [OCR]" if result.ocr_used else ""
                    st.write(f"✓ 解析成功: {result.filename} ({len(result.content)} 字符){ocr_info}")
            except Exception as e:
                item = future_to_item[future]
                failed_files.append({
                    'name': item.get('name', 'unknown'),
                    'original_name': item.get('original_name', item.get('name', 'unknown')),
                    'error': str(e),
                    'size': len(item.get('bytes', b'')),
                    'bytes': item.get('bytes', b'')
                })
                st.write(f"❌ 解析异常: {item.get('name', 'unknown')} - {str(e)}")
            
            if progress_callback:
                progress_callback(i + 1, len(uploaded_items))
    
    st.write(f"📊 解析完成: 成功 {len(parsed_data)}, 失败 {len(failed_files)}")
    return parsed_data, failed_files
            
            if progress_callback:
                progress_callback(i + 1, len(uploaded_items))
    
    st.write(f"📊 解析完成: 成功 {len(parsed_data)}, 失败 {len(failed_files)}")
    return parsed_data, failed_files

# ============================
# 异步API调用 - 提升吞吐量
# ============================
async def call_deepseek_api_async(session: aiohttp.ClientSession, text: str, api_key: str) -> Dict:
    """异步调用DeepSeek API"""
    url = "https://api.siliconflow.cn/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}", 
        "Content-Type": "application/json"
    }
    
    target_city = st.session_state.config.get('target_city', '')
    city_context = f"目标城市：{target_city}" if target_city else ""
    
    system_prompt = f"""你是一个专业的HR简历分析助手。请从简历文本中提取结构化信息。
{city_context}

要求：
1. 严格按照JSON Schema提取字段
2. 学校层次识别：C9 > 985/海外名校 > 211/双一流 > 重点师范/普通一本
3. 评分使用1-5分制
4. 输出必须是合法JSON，不要Markdown格式

JSON Schema:
{JSON_SCHEMA}

缺失字段设为null。"""
    
    # 截断长文本
    truncated_text = text[:12000] if len(text) > 12000 else text
    
    payload = {
        "model": "Pro/deepseek-ai/DeepSeek-V3.2",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"请分析以下简历，提取JSON结构化数据:\n\n{truncated_text}"}
        ],
        "temperature": 0.1,
        "max_tokens": 2500,
        "enable_thinking": False,
        "response_format": {"type": "json_object"}
    }
    
    try:
        timeout = aiohttp.ClientTimeout(total=st.session_state.config.get('api_timeout', 60))
        async with session.post(url, json=payload, headers=headers, timeout=timeout) as response:
            if response.status != 200:
                error_text = await response.text()
                return {"error": f"API错误 {response.status}: {error_text[:200]}"}
            
            result = await response.json()
            if 'choices' not in result or not result['choices']:
                return {"error": "API返回格式异常"}
            
            content = result['choices'][0]['message']['content']
            content = content.replace("```json", "").replace("```", "").strip()
            
            try:
                parsed = json.loads(content)
                # 确保所有必要字段存在
                for key in ['basic_info', 'education', 'work_experience', 'achievements', 'ai_assessment']:
                    if key not in parsed:
                        parsed[key] = {}
                return parsed
            except json.JSONDecodeError as e:
                return {"error": f"JSON解析失败: {str(e)}", "raw_content": content[:500]}
                
    except asyncio.TimeoutError:
        return {"error": "API请求超时"}
    except Exception as e:
        return {"error": f"请求异常: {str(e)}"}

async def process_batch_async(parsed_results: List[ParseResult], api_key: str, progress_callback=None) -> List[Dict]:
    """异步批量处理简历"""
    results = []
    batch_size = st.session_state.config.get('batch_size', 20)
    
    timeout = aiohttp.ClientTimeout(total=st.session_state.config.get('api_timeout', 60))
    
    async with aiohttp.ClientSession(timeout=timeout) as session:
        # 分批处理，控制并发
        for batch_start in range(0, len(parsed_results), batch_size):
            batch = parsed_results[batch_start:batch_start + batch_size]
            
            tasks = []
            for parse_result in batch:
                task = call_deepseek_api_async(session, parse_result.content, api_key)
                tasks.append(task)
            
            # 并发执行当前批次
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            for i, (parse_result, api_result) in enumerate(zip(batch, batch_results)):
                if isinstance(api_result, Exception):
                    api_result = {"error": str(api_result)}
                
                results.append({
                    'filename': parse_result.filename,
                    'parsed_content': parse_result.content[:500] + "..." if len(parse_result.content) > 500 else parse_result.content,
                    'api_result': api_result,
                    'parse_time': parse_result.parse_time
                })
                
                if progress_callback:
                    progress_callback(batch_start + i + 1, len(parsed_results))
            
            # 批次间短暂延迟，避免API限流
            if batch_start + batch_size < len(parsed_results):
                await asyncio.sleep(0.5)
    
    return results

# ============================
# 评分算法 - 基于老版本（3月19日）
# ============================
def calculate_score(data: Dict) -> Tuple[int, str]:
    """
    简历评分算法
    基于2025-03-19版本，通用化设计（移除香港限定）
    """
    score = 0
    logs = []
    
    basic = data.get('basic_info', {})
    edu = data.get('education', {})
    work = data.get('work_experience', {})
    achieve = data.get('achievements', {})
    ai = data.get('ai_assessment', {})
    
    def get_num(val):
        try:
            return float(val)
        except (ValueError, TypeError):
            return 0.0
    
    target_city = st.session_state.config.get('target_city', '')
    
    # 1. 专业匹配 - 教学竞赛
    comp_str = str(achieve.get('teaching_competition', [])) + str(achieve.get('honor_titles', []))
    if "省" in comp_str and ("一等奖" in comp_str or "前三" in comp_str):
        score += 5
        logs.append("专业: 省级奖项 +5")
    
    # 2. 学习经历
    hs_tier = str(edu.get('high_school_tier', ''))
    if "重点" in hs_tier or "县中" in hs_tier:
        score += 3
        logs.append(f"高中: {hs_tier} +3")
    
    b_tier = str(edu.get('bachelor_tier', ''))
    if "C9" in b_tier:
        score += 5
        logs.append("本科: C9 +5")
    elif "985" in b_tier or "海外名校" in b_tier:
        score += 3
        logs.append(f"本科: {b_tier} +3")
    elif "211" in b_tier or "重点师范" in b_tier:
        score += 1
        logs.append(f"本科: {b_tier} +1")
    
    m_tier = str(edu.get('master_tier', ''))
    if "C9" in m_tier:
        score += 5
        logs.append("硕士: C9 +5")
    elif "985" in m_tier or "海外名校" in m_tier:
        score += 3
        logs.append(f"硕士: {m_tier} +3")
    elif "211" in m_tier or "重点师范" in m_tier:
        score += 1
        logs.append(f"硕士: {m_tier} +1")
    
    abroad = get_num(edu.get('study_abroad_years'))
    if abroad >= 2:
        score += 2
        logs.append("留学: 2年以上 +2")
    if str(edu.get('exchange_experience', '否')) == '是':
        score += 1
        logs.append("留学: 交换经历 +1")
    
    # 3. 个人特质 - 基于老版本（男+3，女性已婚已育+1）
    gender = str(basic.get('gender', ''))
    if gender == '男':
        score += 3
        logs.append("性别: 男性 +3")
    if gender == '女' and '已育' in str(basic.get('marital_status', '')):
        score += 1
        logs.append("性别: 已婚已育 +1")
    
    # 家庭背景
    parents = str(basic.get('parents_background', ''))
    if any(k in parents for k in ['教师', '学校', '机关', '研发', '公务员']):
        score += 1
        logs.append("家庭: 父母书香/机关 +1")
    
    # 地域匹配（如果配置了目标城市）
    if target_city:
        if target_city in str(basic.get('residence', '')):
            score += 1
            logs.append(f"地域: 现居{target_city} +1")
        if target_city in str(basic.get('partner_location', '')):
            score += 1
            logs.append(f"地域: 配偶在{target_city} +1")
    
    # 4. 工作经历
    work_tier = str(work.get('school_tier', ''))
    if "重点" in work_tier or "知名" in work_tier or "国际" in work_tier:
        score += 3
        logs.append(f"工作: {work_tier} +3")
    
    if get_num(work.get('non_teaching_gap')) > 2:
        score -= 3
        logs.append("工作: 非教空窗期 -3")
    if get_num(work.get('overseas_work_years')) >= 1:
        score += 3
        logs.append("工作: 海外工作 +3")
    
    # 5. 教学科研
    titles = str(achieve.get('honor_titles', []))
    if any(k in titles for k in ['特级', '学科带头人', '骨干', '优青']):
        score += 5
        logs.append("科研: 核心头衔 +5")
    
    contest = str(achieve.get('teaching_competition', []))
    if "一等奖" in contest and ("区" in contest or "市" in contest or "省" in contest):
        score += 3
        logs.append("科研: 赛课一等奖 +3")
    
    academic = str(achieve.get('academic_results', []))
    if "课题" in academic or "论文" in academic:
        score += 1
        logs.append("科研: 学术成果 +1")
    
    # 6. 管理能力
    mgmt = str(work.get('management_role', ''))
    if mgmt and mgmt not in ['无', '未提及', 'None', 'null']:
        if "年级组长" in mgmt or "教研" in mgmt or "中层" in mgmt or "主任" in mgmt or "校长" in mgmt:
            score += 3
            logs.append(f"管理: {mgmt} +3")
    
    ht_years = get_num(work.get('head_teacher_years'))
    if ht_years >= 5:
        score += 3
        logs.append("管理: 班主任5年+ +3")
    elif ht_years > 0:
        score += 1
        logs.append("管理: 有班主任经历 +1")
    
    # 7. 潜质
    potential = get_num(ai.get('potential_score'))
    if potential > 0:
        score += potential
        logs.append(f"潜质: +{potential}")
    
    return score, "; ".join(logs)

# ============================
# 结果处理与导出
# ============================
def process_results(api_results: List[Dict]) -> List[Dict]:
    """处理API返回结果，计算评分，生成最终数据"""
    final_results = []
    need_review = []
    
    for result in api_results:
        filename = result['filename']
        api_data = result.get('api_result', {})
        
        if 'error' in api_data:
            final_results.append({
                '文件名': filename,
                '处理状态': '失败',
                '错误信息': api_data['error'],
                '综合评分': 0,
                '评分详情': ''
            })
            continue
        
        # 计算评分
        total_score, score_logs = calculate_score(api_data)
        
        # 提取基本信息
        basic = api_data.get('basic_info', {})
        edu = api_data.get('education', {})
        work = api_data.get('work_experience', {})
        achieve = api_data.get('achievements', {})
        ai_eval = api_data.get('ai_assessment', {})
        
        # 检查是否需要人工复核
        needs_review = False
        review_fields = []
        for field_name, key in [('姓名', 'name'), ('性别', 'gender'), ('学科', 'subject')]:
            val = basic.get(key, '')
            if not val or val in ['null', 'None', '']:
                needs_review = True
                review_fields.append(field_name)
        
        row = {
            '文件名': filename,
            '处理状态': '需复核' if needs_review else '成功',
            '姓名': basic.get('name', ''),
            '性别': basic.get('gender', ''),
            '年龄': basic.get('age', ''),
            '任教学科': basic.get('subject', ''),
            '婚姻状况': basic.get('marital_status', ''),
            '现居地': basic.get('residence', ''),
            '本科学校': edu.get('bachelor_school', ''),
            '本科层次': edu.get('bachelor_tier', ''),
            '硕士学校': edu.get('master_school', ''),
            '硕士层次': edu.get('master_tier', ''),
            '现工作单位': work.get('current_company', ''),
            '学校层次': work.get('school_tier', ''),
            '教龄': work.get('teaching_years', ''),
            '班主任年限': work.get('head_teacher_years', ''),
            '管理职务': work.get('management_role', ''),
            '荣誉称号': ', '.join(achieve.get('honor_titles', [])) if isinstance(achieve.get('honor_titles'), list) else str(achieve.get('honor_titles', '')),
            '教学竞赛': ', '.join(achieve.get('teaching_competition', [])) if isinstance(achieve.get('teaching_competition'), list) else str(achieve.get('teaching_competition', '')),
            '综合评分': total_score,
            '评分详情': score_logs,
            'AI评语': ai_eval.get('summary', ''),
            '风险提示': ai_eval.get('risk_warning', ''),
            '需复核字段': ','.join(review_fields) if needs_review else ''
        }
        
        final_results.append(row)
        if needs_review:
            need_review.append(row)
    
    return final_results, need_review

# ============================
# 主界面
# ============================
def main():
    st.title("🎓 智能简历筛选系统")
    st.caption("📊 支持批量解析 | ⚡ 异步处理 | 🎯 智能评分")
    
    # 初始化session state
    for key in ['uploaded_files_queue', 'processing', 'final_results', 'failed_files', 'need_review']:
        if key not in st.session_state:
            st.session_state[key] = [] if 'queue' in key or 'results' in key or 'files' in key else False
    
    # API Key检查
    try:
        api_key = st.secrets["SILICONFLOW_API_KEY"]
    except:
        st.error("🔐 请在 `.streamlit/secrets.toml` 中配置 SILICONFLOW_API_KEY")
        st.code('[secrets]\nSILICONFLOW_API_KEY = "your-api-key"')
        st.stop()
    
    # 侧边栏 - 配置与上传
    with st.sidebar:
        st.header("⚙️ 系统配置")
        
        # 目标城市配置
        target_city = st.text_input(
            "目标城市（可选）",
            value=st.session_state.config.get('target_city', ''),
            placeholder="如：北京、上海，留空则不限制",
            help="配置后会给该城市候选人加分"
        )
        st.session_state.config['target_city'] = target_city
        
        # 性能配置
        col1, col2 = st.columns(2)
        with col1:
            max_workers = st.number_input("并发数", min_value=1, max_value=50, value=10)
            st.session_state.config['max_workers'] = max_workers
        with col2:
            batch_size = st.number_input("批次大小", min_value=1, max_value=50, value=20)
            st.session_state.config['batch_size'] = batch_size
        
        enable_cache = st.checkbox("启用解析缓存", value=True)
        st.session_state.config['enable_cache'] = enable_cache
        
        # 新增：OCR 设置
        st.divider()
        st.subheader("🔍 OCR 设置")
        use_ocr = st.checkbox(
            "启用 OCR（图片型PDF/DOC）",
            value=st.session_state.get('use_ocr', False),
            help="文本提取失败时自动使用 DeepSeek OCR"
        )
        st.session_state['use_ocr'] = use_ocr
        
        st.divider()
        st.header("📁 文件上传")
        st.info("支持: PDF, DOCX, DOC, ZIP, RAR, 7Z")
        
        # 文件上传
        uploaded_files = st.file_uploader(
            "上传简历文件", 
            type=['pdf', 'docx', 'doc'], 
            accept_multiple_files=True,
            help="可一次性选择多个文件"
        )
        
        archive_files = st.file_uploader(
            "上传压缩包",
            type=['zip', 'rar', '7z'],
            accept_multiple_files=True,
            help="自动解压并提取简历文件"
        )
        
        # 处理上传的文件
        new_count = 0
        if uploaded_files:
            existing = [qf['name'] for qf in st.session_state.uploaded_files_queue]
            for f in uploaded_files:
                if f.name not in existing:
                    # 新增：解码文件名
                    decoded_name = decode_filename(f.name)
                    st.session_state.uploaded_files_queue.append({
                        'name': decoded_name,
                        'original_name': f.name,
                        'bytes': f.read()
                    })
                    new_count += 1
        
        if archive_files:
            for archive in archive_files:
                extracted = extract_archive_files(archive.read(), archive.name)
                existing = [qf['name'] for qf in st.session_state.uploaded_files_queue]
                for ef in extracted:
                    if ef['name'] not in existing:
                        st.session_state.uploaded_files_queue.append(ef)
                        new_count += 1
        
        if new_count > 0:
            st.success(f"✅ 新增 {new_count} 个文件，队列共 {len(st.session_state.uploaded_files_queue)} 个")
        
        # 显示队列 - 新增：显示文件名解码对比
        if st.session_state.uploaded_files_queue:
            with st.expander(f"📋 待处理文件 ({len(st.session_state.uploaded_files_queue)} 个)"):
                for f in st.session_state.uploaded_files_queue:
                    original = f.get('original_name', f['name'])
                    decoded = f['name']
                    # 如果文件名被解码了，显示对比
                    if original != decoded:
                        st.text(f"• {decoded}")
                        st.caption(f"  原: {original[:60]}...")
                    else:
                        st.text(f"• {decoded}")
            
            if st.button("🗑️ 清空队列", use_container_width=True):
                st.session_state.uploaded_files_queue = []
                st.rerun()
        
        st.divider()
        
        # 控制按钮 - 始终显示
        st.subheader("▶️ 开始处理")
        
        if not st.session_state.uploaded_files_queue:
            st.warning("⚠️ 请先上传简历文件")
            start_btn = st.button("🚀 开始批量解析", type="primary", use_container_width=True, disabled=True)
        else:
            st.success(f"✅ 已就绪：{len(st.session_state.uploaded_files_queue)} 个文件待处理")
            start_btn = st.button("🚀 开始批量解析", type="primary", use_container_width=True)
        
        # 处理按钮点击 - 必须在 sidebar 内部
        if start_btn and st.session_state.uploaded_files_queue:
            # 保存队列副本并清空
            items_to_process = st.session_state.uploaded_files_queue.copy()
            st.session_state.uploaded_files_queue = []
            
            # 新增：传递 use_ocr 参数
            use_ocr_flag = st.session_state.get('use_ocr', False)
            results = process_all_files(items_to_process, api_key, use_ocr_flag)
            
            # 保存结果到 session_state
            if results['final_results']:
                st.session_state.final_results = results['final_results']
                st.session_state.need_review = results['need_review']
            # 新增：保存失败文件
            if results.get('failed_parse'):
                st.session_state.failed_files = results['failed_parse']
        
        if st.session_state.final_results:
            st.download_button(
                "📥 导出Excel",
                data=pd.DataFrame(st.session_state.final_results).to_excel(index=False),
                file_name=f"简历筛选结果_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                mime="application/vnd.ms-excel",
                use_container_width=True
            )
    
    # 主界面 - 结果显示
    if st.session_state.final_results:
        df = pd.DataFrame(st.session_state.final_results)
        
        # 统计卡片
        st.subheader("📊 处理统计")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("总简历数", len(df))
        with col2:
            success_count = len(df[df['处理状态'] == '成功'])
            st.metric("成功解析", success_count)
        with col3:
            review_count = len(df[df['处理状态'] == '需复核'])
            st.metric("需复核", review_count)
        with col4:
            avg_score = df['综合评分'].mean() if '综合评分' in df.columns else 0
            st.metric("平均评分", f"{avg_score:.1f}")
        
        # 排序显示
        st.subheader("📋 候选人列表")
        if '综合评分' in df.columns:
            df_display = df.sort_values('综合评分', ascending=False)
            st.dataframe(
                df_display.style.background_gradient(subset=['综合评分'], cmap='YlGn'),
                use_container_width=True,
                height=500
            )
        else:
            st.dataframe(df, use_container_width=True)
        
        # 需复核名单
        if st.session_state.need_review:
            st.subheader("⚠️ 需人工复核")
            review_df = pd.DataFrame(st.session_state.need_review)
            st.dataframe(review_df[['文件名', '姓名', '性别', '需复核字段']], use_container_width=True)
        
        # 新增：失败文件下载
        if st.session_state.get('failed_files'):
            st.divider()
            st.subheader("❌ 解析失败的文件")
            failed_list = st.session_state.failed_files
            st.write(f"共 {len(failed_list)} 个文件失败")
            
            # 打包下载
            if st.button("📦 打包下载失败文件"):
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for fail in failed_list:
                        file_bytes = fail.get('bytes', b'')
                        if file_bytes:
                            zf.writestr(fail['name'], file_bytes)
                zip_buffer.seek(0)
                st.download_button(
                    "⬇️ 下载 ZIP",
                    data=zip_buffer.getvalue(),
                    file_name=f"失败文件_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
                    mime="application/zip"
                )
    
# 处理逻辑 - 使用独立的处理函数避免rerun问题
def process_all_files(uploaded_items, api_key, use_ocr=False):
    """处理所有文件的完整流程 - 新增 OCR 支持"""
    total_files = len(uploaded_items)
    results = {
        'parsed_results': [],
        'failed_parse': [],
        'api_results': [],
        'final_results': [],
        'need_review': [],
        'parse_time': 0,
        'ai_time': 0
    }
    
    # 调试信息
    st.info(f"📁 开始处理 {total_files} 个文件...")
    
    try:
        # 阶段1: 文件解析
        with st.spinner(f'📖 正在解析 {total_files} 个文件...'):
            parse_progress = st.progress(0)
            
            def update_parse_progress(current, total):
                parse_progress.progress(current / total)
            
            start_time = time.time()
            # 新增：传递 api_key 和 use_ocr
            parsed_results, failed_parse = parse_files_batch(
                uploaded_items, 
                update_parse_progress,
                api_key=api_key,
                use_ocr=use_ocr
            )
            results['parse_time'] = time.time() - start_time
            results['parsed_results'] = parsed_results
            results['failed_parse'] = failed_parse
            parse_progress.empty()
        
        # 显示解析统计
        st.write(f"📊 解析结果: 成功 {len(parsed_results)}, 失败 {len(failed_parse)}")
        
        if not parsed_results:
            st.error("❌ 没有成功解析的文件")
            return results
        
        st.success(f"✅ 解析完成: {len(parsed_results)} 成功, {len(failed_parse)} 失败 (耗时{results['parse_time']:.1f}s)")
        
        # 阶段2: AI分析（异步）
        with st.spinner(f'🤖 AI正在分析 {len(parsed_results)} 份简历...'):
            ai_progress = st.progress(0)
            
            def update_ai_progress(current, total):
                ai_progress.progress(current / total)
            
            start_time = time.time()
            
            try:
                loop = asyncio.get_event_loop()
            except RuntimeError:
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
            
            api_results = loop.run_until_complete(
                process_batch_async(parsed_results, api_key, update_ai_progress)
            )
            results['ai_time'] = time.time() - start_time
            results['api_results'] = api_results
            ai_progress.empty()
        
        st.success(f"✅ AI分析完成: {len(api_results)} 个 (耗时{results['ai_time']:.1f}s)")
        
        # 阶段3: 结果处理
        with st.spinner('📊 正在生成报告...'):
            final_results, need_review = process_results(api_results)
            results['final_results'] = final_results
            results['need_review'] = need_review
        
        st.success(f"🎉 全部完成！成功: {len([r for r in final_results if r['处理状态'] != '失败'])}, 需复核: {len(need_review)}")
        
    except Exception as e:
        st.error(f"❌ 处理过程中出错: {str(e)}")
        import traceback
        st.code(traceback.format_exc())
    
    return results

if __name__ == "__main__":
    main()
